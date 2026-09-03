"""
Generates "Meridian — Client Discovery & Requirements.docx" — the Phase 2
client-facing companion to MERIDIAN_PHASE2_REPORT.md. Run with:

    python scripts/build_phase2_client_docx.py

Requires python-docx (pip install python-docx).
"""

import datetime
from docx import Document
from docx_helpers import (
    NAVY, GOLD_FILL, GREEN_FILL, GREEN, SLATE, RGBColor,
    setup_document, add_heading, add_body, add_bullets, add_callout,
    add_summary_table, add_questionnaire_table, add_footer, add_cover,
)

OUT_PATH = "Meridian — Client Discovery & Requirements.docx"
TODAY = datetime.date(2026, 9, 3).strftime("%B %d, %Y")

RED = RGBColor(0xB9, 0x1C, 0x1C)
RED_FILL = "FBE4E4"

doc = Document()
setup_document(doc)

# =================================================================== COVER
add_cover(
    doc,
    eyebrow="MERIDIAN — PHASE 2",
    title="Client Discovery & Requirements",
    subtitle="Video rooms, trade broadcasting, and a €250 funded tier — scoped before any of it is built",
    meta_rows=[
        ("Date", TODAY),
        ("Project", "trading-platform (“Meridian”)"),
        ("Purpose", "Confirm Phase 2 scope, gather your answers, flag what needs legal input first"),
    ],
)

# ============================================================ SUMMARY
add_heading(doc, "Quick Summary", level=1)
add_body(
    doc,
    "Meridian v1 (the paper-trading demo) is done and unchanged. Phase 2 proposes three "
    "additions: video consultation rooms, instant trade-broadcast notifications, and a "
    "real-money funded tier starting at a €250 minimum deposit. None of this is built yet "
    "— this document scopes it and collects your answers before we write any code."
)
add_callout(
    doc, "Read this first",
    "Adding real deposits and trade broadcasting changes what kind of business this is. A "
    "demo with fake money is software. A platform that takes real deposits and tells "
    "subscribers what to copy is a financial service — and in the EU, copy-trading has a "
    "specific regulatory classification (see Section 4) that decides whether this needs a "
    "license before launch, not after. This isn't a reason to slow down — it's the reason "
    "the build order in Section 6 puts the legal-sensitive piece last.",
    fill=RED_FILL, border="E8A0A0", label_color=RED,
)

# ============================================================ WHAT'S BUILT
add_heading(doc, "What's Built So Far (v1, unchanged)", level=1)
add_summary_table(
    doc,
    headers=("Feature", "Status"),
    rows=[
        ("3D background canvas", "Done — no changes proposed"),
        ("Paper trading engine", "Done — $100,000 mock balance, 6 instruments, live order book"),
        ("Order execution", "Done — Market/Limit orders, portfolio & activity log"),
        ("Accounts, real data, video, notifications, real money", "Not built — all proposed below"),
    ],
)

# ============================================================ PHASE 2 FEATURES
add_heading(doc, "Phase 2 Additions (Proposed)", level=1)

add_heading(doc, "Video Call & Communication Suite", level=2)
add_body(
    doc,
    "Admin-hosted 1:1 consultations and/or group live-trading rooms, plus a direct-message "
    "support channel. Built on a hosted video provider (not raw WebRTC) — see Section 5 "
    "for the vendor pick. A booking/scheduling flow is a separate piece this needs."
)

add_heading(doc, "Trade Broadcast Notification System", level=2)
add_body(
    doc,
    "The instant an admin executes a trade, subscribed users get an alert — in-app and/or "
    "push notification."
)
add_callout(
    doc, "The most important decision in this document",
    "EU regulators have published explicit guidance on this exact feature: if the platform "
    "automatically places the trade into a follower's account, that is legally “portfolio "
    "management” and requires a MiFID license — regardless of deposit size. If the system "
    "only notifies the user and a person manually places their own order, it does not "
    "trigger that classification. Recommendation: build this as notify-only by design. "
    "Automatic replication is a separate, later decision requiring its own legal review "
    "before any code is written for it.",
    fill=RED_FILL, border="E8A0A0", label_color=RED,
)

add_heading(doc, "Funded Tier — €250 Minimum Deposit", level=2)
add_body(
    doc,
    "A real-money tier alongside the existing free practice tier. Real deposits mean: (1) "
    "KYC becomes mandatory, not optional; (2) funds are routed through a licensed broker "
    "partner rather than held directly — self-custody needs $250,000+ in licensing, a "
    "partner avoids that; (3) since € signals EU/UK users, the broker partner needs EU/UK "
    "regulatory coverage, not just US coverage."
)

# ============================================================ BRANDING
add_heading(doc, "Brand Identity (Retained & Validated)", level=1)
add_summary_table(
    doc,
    headers=("Name", "Trademark Risk"),
    rows=[
        ("Apex Meridian", "None — extends the existing name"),
        ("Voidpoint Capital", "None — original, matches existing design tokens"),
        ("Lodestar Exchange", "None — real historical navigation term"),
    ],
)
add_body(
    doc,
    "Logo shortlist: the Anchor, the Pulse/Sonar Orb, and the Chart + Starburst — all "
    "original marks, zero trademark exposure. (Names and logos with real trademark risk — "
    "GrandLine Capital, LogPose Financial, New World Capital, Pone Markets — were reviewed "
    "and excluded in the prior branding pass.)",
    size=9.5, italic=True, color=SLATE,
)

doc.add_page_break()

# ============================================================ PRICING
add_heading(doc, "Annual Pricing & Financial Infrastructure", level=1)
add_summary_table(
    doc,
    headers=("Service", "Responsible For", "Est. Cost / Year"),
    rows=[
        ("Supabase", "Accounts, trade history, secure state", "$0–300+ (scales with users)"),
        ("Polygon.io", "Unified stock/crypto/forex price feeds", "$0–2,400"),
        ("Daily.co (video, recommended)", "Consultation & live trading rooms", "$0–3,000+ (free to 10,000 min/mo, then $0.004/min)"),
        ("Agora (video, alternative)", "Same, larger-scale alternative", "$0–3,000+ (tiered, from $45.99/mo)"),
        ("Stripe", "Subscriptions and/or €250 deposit processing", "No fixed fee — 2.9% + $0.30/charge"),
        ("Broker partner", "Holding & moving real client funds (funded tier)", "Custom — quote once volume/jurisdiction known"),
    ],
)
add_body(
    doc,
    "Estimates as of this session — confirm current vendor pricing before budgeting.",
    size=9, italic=True, color=SLATE, space_after=4,
)
add_callout(
    doc, "Recommendation & Solution",
    "Supabase + Polygon.io + Daily.co + Stripe covers accounts, data, video, and billing "
    "with zero licensing exposure on their own. The broker partner for the funded tier is "
    "the one line item that's a business decision, not a vendor pick — get a quote and "
    "legal input on EU/UK coverage before committing.",
    fill=GREEN_FILL, border="8FC9A6", label_color=GREEN,
)

doc.add_page_break()

# ============================================================ QUESTIONNAIRE
add_heading(doc, "Updated Stakeholder Discovery Questionnaire", level=1)
add_body(doc, "Type your answer directly into the shaded box next to each question.")

add_heading(doc, "1. Trading Basics & Asset Classes", level=2)
add_questionnaire_table(doc, [
    (1, "Which asset classes are in scope for the funded tier?", "Stocks, crypto, forex — or a mix."),
    (2, "Does the €250 minimum apply platform-wide, or only to specific tiers/assets?", ""),
    (3, "Is €250 a one-time minimum to open the tier, or also a minimum per top-up?", ""),
    (4, "Must practice-tier users graduate through a step before the funded tier?", "Or can they join the funded tier directly?"),
])

add_heading(doc, "2. Video & Communication Rules", level=2)
add_questionnaire_table(doc, [
    (5, "One-on-one consultations, group live-trading rooms, or both?", ""),
    (6, "Who can host a room — only the admin, or designated staff too?", ""),
    (7, "Is scheduling self-serve, or admin-invited only?", ""),
    (8, "Should sessions be recorded?", "If yes: who can access recordings, and for how long?"),
])

add_heading(doc, "3. Trade Broadcast Notifications", level=2)
add_questionnaire_table(doc, [
    (9, "Notify-only, or automatic replication into follower accounts?", "This answer decides whether MiFID licensing is required — a legal decision, not a preference."),
    (10, "Which channels — in-app, push, email, SMS, or a mix?", ""),
    (11, "Does every trade broadcast, or only ones the admin flags for sharing?", ""),
    (12, "Do notifications include rationale/commentary, or just the raw fill?", ""),
])

add_heading(doc, "4. Look, Branding & Data Security", level=2)
add_questionnaire_table(doc, [
    (13, "Confirm final naming: Apex Meridian, Voidpoint Capital, or Lodestar Exchange?", ""),
    (14, "Confirm final logo: the anchor, the pulse/sonar orb, or the chart + starburst?", ""),
    (15, "EU users are now in scope (the € deposit) — who is the GDPR data controller, and where is data stored?", ""),
    (16, "Any jurisdictions beyond the EU/UK planned for the funded tier?", ""),
])

# ============================================================ NEXT STEPS
add_heading(doc, "Recommended Build Sequence", level=1)
add_bullets(doc, [
    "1. Accounts & persistence (Supabase) — needed before everything else here.",
    "2. Real market data (Polygon.io) — replaces the client-side random walk.",
    "3. Notify-only trade broadcasting — explicitly not auto-replicating trades.",
    "4. Video consulting (Daily.co) — 1:1 rooms first, group rooms once validated.",
    "5. Funded tier (€250 minimum) — only after a broker partner with EU/UK coverage is "
    "contracted, KYC is implemented, and legal has confirmed Question 9's answer doesn't "
    "trigger portfolio-management licensing.",
])
add_body(
    doc,
    "Steps 1–4 are engineering-led. Step 5 is legal-led with engineering support — it "
    "should not start until Section 3, Question 9 has a confirmed answer.",
    italic=True, size=9.5, color=SLATE,
)

add_footer(doc, "Meridian — Phase 2 Client Discovery & Requirements")

doc.save(OUT_PATH)
print(f"Saved: {OUT_PATH}")
