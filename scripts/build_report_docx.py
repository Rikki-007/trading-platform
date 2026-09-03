"""
Generates platform-report.docx — a short, plain-language Word version of
PLATFORM_REPORT.md. Run with:

    python scripts/build_report_docx.py

Requires python-docx (pip install python-docx).
"""

import datetime
from docx import Document
from docx_helpers import (
    NAVY, GOLD, GOLD_FILL, GREEN, GREEN_FILL, SLATE,
    setup_document, add_heading, add_body, add_bullets, add_callout,
    add_summary_table, add_questionnaire_table, add_footer, add_cover,
)

OUT_PATH = "platform-report.docx"
TODAY = datetime.date(2026, 9, 3).strftime("%B %d, %Y")

doc = Document()
setup_document(doc)

# =================================================================== COVER
add_cover(
    doc,
    eyebrow="MERIDIAN",
    title="Platform Report",
    subtitle="What we built, what's next, and the questions we need answered",
    meta_rows=[
        ("Date", TODAY),
        ("Project", "trading-platform (“Meridian”)"),
        ("Status", "v1 — working demo, paper trading only"),
    ],
)

# ============================================================ WHAT THIS IS
add_heading(doc, "What This Is", level=1)
add_body(
    doc,
    "Meridian is a working demo of a trading app. It looks and feels real — "
    "prices move, you can buy and sell, your balance updates — but it's all "
    "practice money. Nothing here connects to real markets, real accounts, or "
    "real money yet."
)

# ============================================================ WHAT'S BUILT
add_heading(doc, "What's Built Right Now", level=1)

add_heading(doc, "3D background", level=3, space_before=10)
add_bullets(doc, [
    "A moving 3D scene behind the app — glowing rings and floating particles.",
    "It reacts to your mouse and to scrolling, and the colors shift as you scroll down.",
])

add_heading(doc, "Trading dashboard", level=3)
add_bullets(doc, [
    "6 practice stocks with prices that update every couple seconds.",
    "$100,000 in practice money to start.",
    "A buy/sell screen with Market and Limit order types.",
    "A live order book (shows fake buy/sell offers around the current price).",
    "A portfolio screen showing your cash, holdings, and gains/losses.",
    "A history log of every trade you make.",
])

add_heading(doc, "Look and feel", level=3)
add_bullets(doc, [
    "Dark theme with cyan, gold, and red accent colors.",
    "Glass-style navigation bar and side menu.",
    "Smooth animations throughout (Framer Motion).",
])

add_heading(doc, "Not Built Yet", level=2)
add_body(
    doc,
    "To be upfront: a scrolling price ticker and a “system status” indicator were "
    "mentioned early on but aren't in the app yet. Easy to add later — just flagging "
    "it now so nothing is assumed.",
    italic=True, color=SLATE, size=9.5,
)
add_bullets(doc, [
    "No real accounts or login — refreshing the page resets everything.",
    "No real market data — all prices are made up.",
    "No real money, no payments, no ID checks.",
    "No saved history — nothing is stored anywhere.",
])

doc.add_page_break()

# ============================================================ QUESTIONS
add_heading(doc, "Questions For You", level=1)
add_body(
    doc, "Please answer these so we can plan what to build next. Just type your "
         "answer in the shaded box next to each question."
)

add_heading(doc, "1. Trading basics", level=2)
add_questionnaire_table(doc, [
    (1, "What should people trade — stocks, options, crypto, forex?", "Changes which data/broker we need."),
    (2, "Stay practice-only, or add real money later?", "Real money = broker partnership + legal work."),
    (3, "What order types do we need beyond Buy/Sell (Market/Limit)?", "e.g. stop-loss, trailing stop."),
    (4, "Should users be able to short-sell or use leverage?", "Not in the app today."),
    (5, "Who is this for — beginners, active traders, a specific group?", ""),
])

add_heading(doc, "2. Account rules", level=2)
add_questionnaire_table(doc, [
    (6, "Should the starting balance always be $100,000, or should it vary?", ""),
    (7, "Do we need margin (borrowing to trade)?", "If yes, what are the limits?"),
    (8, "Should there be daily loss limits?", "More useful for teaching, less for pure practice."),
    (9, "Do you want leaderboards or competitions?", ""),
])

add_heading(doc, "3. Look and branding", level=2)
add_questionnaire_table(doc, [
    (10, "Does the current dark/glowing style feel right?", "Or should it be more plain/corporate, or more playful?"),
    (11, "Do you have an existing logo or brand to use?", "Or do we keep building on “Meridian”?"),
    (12, "Any language or accessibility needs?", "Easier to plan now than to change later."),
])

add_heading(doc, "4. Accounts & data", level=2)
add_questionnaire_table(doc, [
    (13, "How should people log in — email, Google, Apple, wallet?", ""),
    (14, "Should trade history be kept forever, or just for a while?", ""),
    (15, "Does user data need to stay in a specific country?", ""),
    (16, "Should progress sync across devices (phone + laptop)?", "This needs a real server."),
])

doc.add_page_break()

# ============================================================ WHAT'S NEXT
add_heading(doc, "What's Next: Your Options", level=1)
add_body(
    doc, "Here are the next 3 things to build, in order. For each one, here are "
         "the choices and what we'd recommend."
)

add_heading(doc, "1. Real market prices", level=2)
add_summary_table(
    doc,
    headers=("Option", "What it's good for"),
    rows=[
        ("Polygon.io", "Covers stocks, crypto, and forex in one place. Good all-around choice."),
        ("Alpaca", "Good if we also use Alpaca to handle real trades later."),
        ("Alpha Vantage", "Free to test with, but too limited once real users show up."),
        ("Binance", "Only for crypto prices."),
    ],
)
add_callout(doc, "Recommendation", "Start with Polygon.io. One vendor, covers everything we need.",
            fill=GREEN_FILL, border="8FC9A6", label_color=GREEN)

add_heading(doc, "2. Real accounts & login", level=2)
add_summary_table(
    doc,
    headers=("Option", "What it's good for"),
    rows=[
        ("Supabase", "Login + database in one. Fits trading data well (accounts, trades, positions)."),
        ("Firebase", "Similar, but fits this kind of data less naturally."),
        ("Build it ourselves", "Full control, but much more work for no real benefit here."),
    ],
)
add_callout(doc, "Recommendation", "Use Supabase. It gives us login and a database together, "
                                    "and it fits how trades and positions naturally work.",
            fill=GREEN_FILL, border="8FC9A6", label_color=GREEN)

add_heading(doc, "3. Payments", level=2)
add_summary_table(
    doc,
    headers=("Use case", "What it's good for"),
    rows=[
        ("Paid subscriptions (no real trading)", "Stripe. Simple, standard, no legal complications."),
        ("Real-money stock trading", "Partner with a licensed broker (e.g. Alpaca) — they hold the money, not us."),
        ("Real-money crypto trading", "Partner with a licensed crypto service — same idea, they hold the funds."),
        ("Holding money ourselves", "Not recommended — this requires special licenses."),
    ],
)
add_callout(doc, "Recommendation",
            "If you want a paid tier with no real trading, use Stripe — easy. If you want real "
            "trading later, partner with a licensed broker instead of holding funds ourselves — "
            "much simpler and avoids needing special licenses.",
            fill=GREEN_FILL, border="8FC9A6", label_color=GREEN)

add_heading(doc, "Suggested order to build things", level=2)
add_bullets(doc, [
    "1. Accounts & login (Supabase)",
    "2. Real market prices (Polygon.io)",
    "3. Paid subscriptions, if wanted (Stripe)",
    "4. Real-money trading — only after you confirm you want it, and only with legal advice first",
])

add_footer(doc, "Meridian — Platform Report")

doc.save(OUT_PATH)
print(f"Saved: {OUT_PATH}")
