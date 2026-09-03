"""
Shared styling helpers for the project's python-docx generator scripts.
Keeps both build_report_docx.py and build_client_docx.py using the same
palette, fonts, and table/heading building blocks instead of duplicating
oxml boilerplate.
"""

from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ---------------------------------------------------------------- palette --
NAVY = RGBColor(0x16, 0x1B, 0x2C)
NAVY_FILL = "16213E"
GOLD = RGBColor(0xB8, 0x86, 0x0B)
GOLD_FILL = "F3E6CC"
GREEN = RGBColor(0x1F, 0x7A, 0x4D)
GREEN_FILL = "E3F1E9"
SLATE = RGBColor(0x47, 0x55, 0x69)
INK = RGBColor(0x22, 0x27, 0x33)
LINE_GREY = "D9DEE6"
ROW_ALT_FILL = "F5F7FA"
ANSWER_FILL = "FFF7DE"
WHITE = RGBColor(0xFF, 0xFF, 0xFF)

FONT_BODY = "Calibri"
FONT_HEAD = "Calibri"


def set_cell_shading(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, color=LINE_GREY, size=4):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(size))
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)
        borders.append(el)
    tcPr.append(borders)


def set_cell_margins(cell, top=60, bottom=60, left=100, right=100):
    tcPr = cell._tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for edge, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        node = OxmlElement(f"w:{edge}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        mar.append(node)
    tcPr.append(mar)


def set_repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    el = OxmlElement("w:tblHeader")
    el.set(qn("w:val"), "true")
    trPr.append(el)


def add_bottom_border(paragraph, color=LINE_GREY, size=8, space=6):
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    pPr.append(pbdr)


def set_cell_text(cell, text, *, bold=False, italic=False, color=INK, size=10,
                   font=FONT_BODY, align=None, shading=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = font
    run.font.color.rgb = color
    set_cell_borders(cell)
    set_cell_margins(cell)
    if shading:
        set_cell_shading(cell, shading)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


def setup_document(doc):
    normal = doc.styles["Normal"]
    normal.font.name = FONT_BODY
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.18

    section = doc.sections[0]
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    return section


def add_heading(doc, text, level=1, space_before=18, space_after=8, rule=True):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = FONT_HEAD
    run.bold = True
    if level == 1:
        run.font.size = Pt(17)
        run.font.color.rgb = NAVY
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = NAVY
    else:
        run.font.size = Pt(11.5)
        run.font.color.rgb = GOLD
    if rule and level == 1:
        add_bottom_border(p)
    return p


def add_body(doc, text, *, size=10.5, color=INK, italic=False, bold=False, space_after=8):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.italic = italic
    run.bold = bold
    return p


def add_bullets(doc, items, *, size=10.5):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(item)
        run.font.size = Pt(size)
        run.font.color.rgb = INK


def add_callout(doc, label, text, fill=GOLD_FILL, border="E3C77A", label_color=GOLD):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    set_cell_shading(cell, fill)
    set_cell_borders(cell, color=border, size=4)
    set_cell_margins(cell, top=140, bottom=140, left=180, right=180)
    cell.text = ""
    p1 = cell.paragraphs[0]
    r1 = p1.add_run(label.upper())
    r1.bold = True
    r1.font.size = Pt(8.5)
    r1.font.color.rgb = label_color
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(2)
    r2 = p2.add_run(text)
    r2.font.size = Pt(10)
    r2.font.color.rgb = NAVY
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return tbl


def add_summary_table(doc, rows, headers=("Area", "Status"), widths=None):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = True
    hdr = tbl.rows[0]
    set_repeat_header(hdr)
    for i, h in enumerate(headers):
        set_cell_text(hdr.cells[i], h, bold=True, color=WHITE, size=9.5,
                      shading=NAVY_FILL, align=WD_ALIGN_PARAGRAPH.LEFT)
    for r_i, row in enumerate(rows):
        tr = tbl.add_row()
        fill = ROW_ALT_FILL if r_i % 2 == 1 else None
        for i, val in enumerate(row):
            set_cell_text(tr.cells[i], val, bold=(i == 0), size=10, shading=fill)
    if widths:
        for row in tbl.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return tbl


def add_questionnaire_table(doc, rows):
    """rows: list of (n, question, guidance) — adds a Client Response column."""
    headers = ("#", "Question", "Guidance", "Client Response")
    widths = (0.35, 2.55, 2.15, 1.45)
    tbl = doc.add_table(rows=1, cols=4)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tbl.rows[0]
    set_repeat_header(hdr)
    for i, h in enumerate(headers):
        set_cell_text(hdr.cells[i], h, bold=True, color=WHITE, size=9, shading=NAVY_FILL,
                      align=WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT)
    for n, question, guidance in rows:
        tr = tbl.add_row()
        set_cell_text(tr.cells[0], str(n), size=9.5, align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(tr.cells[1], question, bold=True, size=9.5)
        set_cell_text(tr.cells[2], guidance, italic=True, color=SLATE, size=8.5)
        set_cell_text(tr.cells[3], "", shading=ANSWER_FILL, size=9.5)
    for row in tbl.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return tbl


def add_footer(doc, text):
    section = doc.sections[0]
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = fp.add_run(text)
    run.font.size = Pt(8)
    run.font.color.rgb = SLATE


def add_cover(doc, eyebrow, title, subtitle, meta_rows):
    p0 = doc.add_paragraph()
    p0.paragraph_format.space_before = Pt(150)
    r0 = p0.add_run(eyebrow)
    r0.font.name = FONT_HEAD
    r0.font.size = Pt(15)
    r0.font.color.rgb = GOLD
    r0.bold = True

    p1 = doc.add_paragraph()
    p1.paragraph_format.space_before = Pt(4)
    p1.paragraph_format.space_after = Pt(2)
    r1 = p1.add_run(title)
    r1.font.name = FONT_HEAD
    r1.font.size = Pt(30)
    r1.font.color.rgb = NAVY
    r1.bold = True

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(220)
    r2 = p2.add_run(subtitle)
    r2.font.size = Pt(13)
    r2.font.color.rgb = SLATE
    r2.italic = True

    meta = doc.add_table(rows=len(meta_rows), cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.LEFT
    for i, (k, v) in enumerate(meta_rows):
        set_cell_text(meta.rows[i].cells[0], k, bold=True, size=10, color=SLATE)
        set_cell_text(meta.rows[i].cells[1], v, size=10)
    for row in meta.rows:
        row.cells[0].width = Inches(1.6)
        row.cells[1].width = Inches(4.9)

    doc.add_page_break()
