"""
Generates platform-client-discovery-and-requirements.docx — a short,
client-facing companion to platform-report.docx. Run with:

    python scripts/build_client_docx.py

Requires python-docx (pip install python-docx).
"""

import datetime
from docx import Document
from docx_helpers import (
    GREEN_FILL, GREEN, SLATE,
    setup_document, add_heading, add_body, add_bullets, add_callout,
    add_summary_table, add_questionnaire_table, add_footer, add_cover,
)

OUT_PATH = "platform-client-discovery-and-requirements.docx"
TODAY = datetime.date(2026, 9, 3).strftime("%B %d, %Y")

doc = Document()
setup_document(doc)

# =================================================================== COVER
add_cover(
    doc,
    eyebrow="MERIDIAN",
    title="Client Discovery & Requirements",
    subtitle="A short summary of what's built, plus questions we need answered",
    meta_rows=[
        ("Date", TODAY),
        ("Project", "trading-platform (“Meridian”)"),
        ("Purpose", "Confirm what's built, gather your answers, plan next steps"),
    ],
)

# ============================================================ SUMMARY
add_heading(doc, "Quick Summary", level=1)
add_body(
    doc,
    "We've built a working demo of a trading app. It's fully practice-only — no real "
    "market data, no real accounts, no real money. Everything resets when you refresh "
    "the page. Before building the next phase (real data, real accounts, maybe real "
    "money), we need your answers to a few questions — see below."
)

# ============================================================ WHAT'S BUILT
add_heading(doc, "What's Built So Far", level=1)
add_summary_table(
    doc,
    headers=("Feature", "What it does"),
    rows=[
        ("3D background", "A moving, glowing 3D scene that reacts to your mouse and scrolling."),
        ("Trading dashboard", "6 practice stocks with live-moving prices, starting with $100,000 in fake money."),
        ("Buy & sell", "A working order form (Market and Limit orders) that actually updates your balance."),
        ("Order book", "A live list of fake buy/sell offers around the current price."),
        ("Portfolio & history", "Shows your cash, holdings, gains/losses, and every trade you've made."),
        ("Design", "Dark theme, cyan/gold/red accents, smooth animations throughout."),
    ],
)
add_body(
    doc,
    "Not built yet: a scrolling price ticker and a “system status” light were mentioned "
    "early on, but they're not in the app. Simple to add later — just being upfront.",
    italic=True, color=SLATE, size=9.5,
)

doc.add_page_break()

# ============================================================ QUESTIONNAIRE
add_heading(doc, "Questions For You", level=1)
add_body(doc, "Type your answer directly into the shaded box next to each question.")

add_heading(doc, "1. Trading Basics", level=2)
add_questionnaire_table(doc, [
    (1, "What should people trade — stocks, options, crypto, forex?", "Changes what data/broker we need."),
    (2, "Stay practice-only, or add real money later?", "Real money needs a broker partner and legal review."),
    (3, "What order types beyond Buy/Sell (Market/Limit)?", "e.g. stop-loss, trailing stop."),
    (4, "Should users be able to short-sell or use leverage?", "Not in the app today."),
    (5, "Who is this for?", "e.g. beginners, active traders, a specific group."),
])

add_heading(doc, "2. Account Rules", level=2)
add_questionnaire_table(doc, [
    (6, "Should everyone start with $100,000, or should it vary?", ""),
    (7, "Do we need margin (borrowing to trade)?", "If yes, what are the limits?"),
    (8, "Should there be daily loss limits?", "Useful for teaching, less so for pure practice."),
    (9, "Do you want leaderboards or competitions?", ""),
])

add_heading(doc, "3. Look & Branding", level=2)
add_questionnaire_table(doc, [
    (10, "Does the current dark/glowing style feel right?", "Or more plain/corporate, or more playful?"),
    (11, "Do you have an existing logo or brand to use?", "Or keep building on “Meridian”?"),
    (12, "Any language or accessibility needs?", "Easier to plan now than change later."),
])

add_heading(doc, "4. Accounts & Data", level=2)
add_questionnaire_table(doc, [
    (13, "How should people log in?", "e.g. email, Google, Apple, wallet."),
    (14, "Keep trade history forever, or just for a while?", ""),
    (15, "Does user data need to stay in a specific country?", ""),
    (16, "Should progress sync across devices?", "Needs a real server, not just the browser."),
])

doc.add_page_break()

# ============================================================ ROADMAP
add_heading(doc, "What's Next: Options & Recommendations", level=1)
add_body(doc, "Three things to build next, in order. For each, here are the choices and our pick.")

add_heading(doc, "1. Real market prices", level=2)
add_summary_table(
    doc,
    headers=("Option", "What it's good for"),
    rows=[
        ("Polygon.io", "Covers stocks, crypto, and forex in one place."),
        ("Alpaca", "Good if we also use Alpaca for real trades later."),
        ("Alpha Vantage", "Free to test with, too limited once real users show up."),
        ("Binance", "Crypto prices only."),
    ],
)
add_callout(doc, "Recommendation", "Polygon.io — one vendor, covers everything we need.",
            fill=GREEN_FILL, border="8FC9A6", label_color=GREEN)

add_heading(doc, "2. Real accounts & login", level=2)
add_summary_table(
    doc,
    headers=("Option", "What it's good for"),
    rows=[
        ("Supabase", "Login + database together. Fits trading data well."),
        ("Firebase", "Similar, but fits this kind of data less naturally."),
        ("Build it ourselves", "Full control, much more work, no real benefit here."),
    ],
)
add_callout(doc, "Recommendation", "Supabase — login and database in one, fits how trades and "
                                    "positions naturally work.",
            fill=GREEN_FILL, border="8FC9A6", label_color=GREEN)

add_heading(doc, "3. Payments", level=2)
add_summary_table(
    doc,
    headers=("Use case", "What it's good for"),
    rows=[
        ("Paid subscriptions (no real trading)", "Stripe — simple, standard, no legal complications."),
        ("Real-money stock trading", "Partner with a licensed broker — they hold the money, not us."),
        ("Real-money crypto trading", "Partner with a licensed crypto service — same idea."),
        ("Holding money ourselves", "Not recommended — needs special licenses."),
    ],
)
add_callout(
    doc, "Recommendation",
    "For a paid tier with no real trading: Stripe. For real trading later: partner with a "
    "licensed broker instead of holding funds ourselves — simpler, and avoids needing "
    "special licenses.",
    fill=GREEN_FILL, border="8FC9A6", label_color=GREEN,
)

# ============================================================ NEXT STEPS
add_heading(doc, "Next Steps", level=1)
add_bullets(doc, [
    "Fill in the questions above and send this back to us.",
    "We'll confirm the plan based on your answers.",
    "If you want real-money trading, we'll set up a separate legal/compliance conversation "
    "before building that part.",
])

add_footer(doc, "Meridian — Client Discovery & Requirements")

doc.save(OUT_PATH)
print(f"Saved: {OUT_PATH}")
